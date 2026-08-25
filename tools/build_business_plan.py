from pathlib import Path
import re
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "AI_SWARMER_OS_COMPANY_BUILD_AND_GTM_PLAN.md"
OUTPUT = ROOT / "deliverables" / "AI_SWARMER_OS_Company_Build_and_GTM_Plan.docx"

NAVY = "07111A"
PANEL = "10202A"
TEAL = "35D6B4"
BLUE = "69A7FF"
AMBER = "F0B65A"
WHITE = "F2F7F8"
INK = "15232C"
MUTED = "60717C"
LINE = "C9D5D9"
PALE = "EAF4F2"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = "w:" + edge
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            for key in ("val", "sz", "space", "color"):
                if key in kwargs[edge]:
                    element.set(qn("w:" + key), str(kwargs[edge][key]))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("AI SWARMER OS  •  ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def set_repeat_header_text(section):
    p = section.header.paragraphs[0]
    p.text = "VERLORAY SECURITY TECHNOLOGY INNOVATIONS     /     WILKERSON COLLECTIVE"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MUTED)
        run.font.letter_spacing = Pt(0.8)


def add_field_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Open in Word and update this field to refresh the table of contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def strip_inline(text):
    return re.sub(r"\*\*(.*?)\*\*", r"\1", text.replace("`", ""))


def add_rich_text(paragraph, text):
    # Minimal inline Markdown parser for bold and inline code.
    token = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    pos = 0
    for match in token.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        value = match.group(0)
        if value.startswith("**"):
            run = paragraph.add_run(value[2:-2])
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(INK)
        else:
            run = paragraph.add_run(value[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string("116F65")
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in [
        ("Title", 34, NAVY, 0, 14),
        ("Heading 1", 21, NAVY, 20, 8),
        ("Heading 2", 13.5, "137A6C", 14, 5),
        ("Heading 3", 10.5, NAVY, 10, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    quote = doc.styles["Quote"]
    quote.font.name = "Aptos"
    quote.font.size = Pt(9.2)
    quote.font.italic = False
    quote.font.color.rgb = RGBColor.from_string("315D57")
    quote.paragraph_format.left_indent = Inches(0.25)
    quote.paragraph_format.right_indent = Inches(0.15)
    quote.paragraph_format.space_before = Pt(7)
    quote.paragraph_format.space_after = Pt(7)

    for style_name in ("List Bullet", "List Number"):
        s = doc.styles[style_name]
        s.font.name = "Aptos"
        s.font.size = Pt(9.2)
        s.font.color.rgb = RGBColor.from_string(INK)
        s.paragraph_format.left_indent = Inches(0.28)
        s.paragraph_format.first_line_indent = Inches(-0.14)
        s.paragraph_format.space_after = Pt(2.8)

    if "Checklist" not in doc.styles:
        checklist = doc.styles.add_style("Checklist", WD_STYLE_TYPE.PARAGRAPH)
        checklist.base_style = doc.styles["Normal"]
        checklist.font.name = "Aptos"
        checklist.font.size = Pt(9.2)
        checklist.paragraph_format.left_indent = Inches(0.28)
        checklist.paragraph_format.first_line_indent = Inches(-0.16)


def add_cover(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, NAVY)
    set_cell_border(cell, top={"val":"nil"}, left={"val":"nil"}, bottom={"val":"nil"}, right={"val":"nil"})
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.width = Inches(6.9)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("AI SWARMER OS")
    r.font.name = "Aptos Display"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(54)
    p2.paragraph_format.space_after = Pt(12)
    r = p2.add_run("Company Build, Funding,\nBrand & Go-to-Market Plan")
    r.font.name = "Aptos Display"
    r.font.size = Pt(31)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(WHITE)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("A 36-month operating system for a defensible AI-agent security company")
    r.font.name = "Aptos"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("A9C8C2")
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(58)
    p4.paragraph_format.space_after = Pt(22)
    r = p4.add_run("VERLORAY SECURITY TECHNOLOGY INNOVATIONS\nWILKERSON COLLECTIVE\n\nVERSION 1.0  /  03 AUGUST 2026")
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("CONFIDENTIAL STRATEGY DOCUMENT")
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Truthful claims • Secure-by-design • Evidence before scale")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("137A6C")
    doc.add_page_break()


def add_navigation(doc):
    p = doc.add_paragraph("Executive navigation", style="Title")
    p.runs[0].font.size = Pt(27)
    intro = doc.add_paragraph("This plan is designed as a maintained operating document. Update the Word table below after edits, and record material terminology or strategy changes in the Founder’s Brain source file.")
    intro.paragraph_format.space_after = Pt(16)
    toc = doc.add_paragraph()
    add_field_toc(toc)
    note = doc.add_paragraph(style="Quote")
    note.add_run("Decision rule: ").bold = True
    note.add_run("advance claims from Roadmap to Pilot to Working only when a test, customer acceptance, or independently reviewable artifact supports the change.")
    doc.add_page_break()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    widths = [Inches(6.95 / len(headers))] * len(headers)
    if len(headers) == 4:
        widths = [Inches(1.35), Inches(2.8), Inches(1.15), Inches(1.65)]
    elif len(headers) == 3:
        widths = [Inches(1.55), Inches(3.05), Inches(2.35)]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        shade(cell, PANEL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(strip_inline(text))
        r.font.name = "Aptos"
        r.font.size = Pt(7.4)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.width = widths[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if ridx % 2:
                shade(cell, "F3F7F7")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_rich_text(p, text)
            for run in p.runs:
                run.font.size = Pt(7.8)
        for cell in row.cells:
            set_cell_border(cell,
                top={"val":"single","sz":"4","color":LINE},
                left={"val":"single","sz":"4","color":LINE},
                bottom={"val":"single","sz":"4","color":LINE},
                right={"val":"single","sz":"4","color":LINE})
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_markdown(doc, text):
    lines = text.splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## 1. Executive decision"))
    i = start
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith("| "):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            parsed = [[strip_inline(c.strip()) for c in row.strip("|").split("|")] for row in block]
            if len(parsed) >= 2:
                headers = parsed[0]
                rows = parsed[2:] if all(re.fullmatch(r":?-{3,}:?", c) for c in parsed[1]) else parsed[1:]
                add_table(doc, headers, rows)
            continue
        if line.startswith("### "):
            doc.add_paragraph(strip_inline(line[4:]), style="Heading 2")
        elif line.startswith("## "):
            p = doc.add_paragraph(strip_inline(line[3:]), style="Heading 1")
            p.paragraph_format.page_break_before = i != start
        elif line.startswith("#### "):
            doc.add_paragraph(strip_inline(line[5:]), style="Heading 3")
        elif line.startswith("> "):
            p = doc.add_paragraph(style="Quote")
            add_rich_text(p, line[2:])
        elif re.match(r"^- \[[ xX]\] ", line):
            checked = line[3].lower() == "x"
            p = doc.add_paragraph(style="Checklist")
            add_rich_text(p, ("☒ " if checked else "☐ ") + line[6:])
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\. ", "", line))
        else:
            paragraph_lines = [line]
            while i + 1 < len(lines):
                nxt = lines[i + 1].strip()
                if not nxt or nxt.startswith(("#", "- ", "> ", "|")) or re.match(r"^\d+\. ", nxt):
                    break
                paragraph_lines.append(nxt)
                i += 1
            p = doc.add_paragraph()
            add_rich_text(p, " ".join(paragraph_lines))
        i += 1


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    set_repeat_header_text(doc.sections[0])
    add_page_number(doc.sections[0].footer.paragraphs[0])
    add_cover(doc)
    add_navigation(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    props = doc.core_properties
    props.title = "AI SWARMER OS — Company Build, Funding, Brand & Go-to-Market Plan"
    props.subject = "36-month company and go-to-market operating plan"
    props.author = "Wilkerson Collective / Verloray Security Technology Innovations"
    props.keywords = "AI SWARMER OS, cybersecurity, AI agents, go-to-market, business plan"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

